# -*- coding: utf-8 -*-
#
#
# Project name: OpenVAS Reporting: A tool to convert OpenVAS XML reports into Excel files.
# Project URL: https://github.com/TheGroundZero/openvas_to_report

import re
from collections import Counter

from .config import Config
from .parsed_data import Vulnerability


def exporters():
    """
    Enum-like instance containing references to correct exporter function

    > exporters()[key](param[s])

    :return: Pointer to exporter function
    """
    return {
        'xlsx': export_to_excel,
        'docx': export_to_word,
        'csv': export_to_csv
    }


def export_to_excel(vuln_info, template=None, output_file='openvas_report.xlsx'):
    """
    Export vulnerabilities info in an Excel file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)
    :param template: Not supported in xlsx-output
    :type template: NoneType

    :param output_file: Filename of the Excel file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import xlsxwriter

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in XSLX-output.")

    # ====================
    # FUNCTIONS
    # ====================
    def __row_height(text, width):
        return (max((len(text) // width), text.count('\n')) + 1) * 15

    workbook = xlsxwriter.Workbook(output_file)

    workbook.set_properties({
        'title': output_file,
        'subject': 'OpenVAS report',
        'author': 'TheGroundZero',
        'category': 'report',
        'keywords': 'OpenVAS, report',
        'comments': 'TheGroundZero (https://github.com/TheGroundZero)'})

    # ====================
    # FORMATTING
    # ====================
    workbook.formats[0].set_font_name('Tahoma')

    format_sheet_title_content = workbook.add_format({'font_name': 'Tahoma', 'font_size': 12,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'center', 'valign': 'vcenter', 'border': 1})
    format_table_titles = workbook.add_format({'font_name': 'Tahoma', 'font_size': 11,
                                               'font_color': 'white', 'bold': True,
                                               'align': 'center', 'valign': 'vcenter',
                                               'border': 1,
                                               'bg_color': Config.colors()['blue']})
    format_table_cells = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                              'align': 'left', 'valign': 'top',
                                              'border': 1, 'text_wrap': 1})
    format_align_center = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top'})
    format_align_border = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_toc = {
        'critical': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                         'align': 'center', 'valign': 'top',
                                         'border': 1,
                                         'bg_color': Config.colors()['critical']}),
        'high': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['high']}),
        'medium': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                       'align': 'center', 'valign': 'top',
                                       'border': 1, 'bg_color': Config.colors()['medium']}),
        'low': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                    'align': 'center', 'valign': 'top',
                                    'border': 1, 'bg_color': Config.colors()['low']}),
        'none': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['none']})
    }

    # TODO Move to function to de-duplicate this
    vuln_info.sort(key=lambda key: key.cvss, reverse=True)
    vuln_levels = Counter()
    vuln_host_by_level = Counter()
    vuln_by_family = Counter()

    for i, vuln in enumerate(vuln_info, 1):
        vuln_levels[vuln.level.lower()] += 1
        vuln_host_by_level[vuln.level.lower()] += len(vuln.hosts)
        vuln_by_family[vuln.family] += 1

    # ====================
    # SUMMARY SHEET
    # ====================
    sheet_name = "Summary"
    ws_sum = workbook.add_worksheet(sheet_name)
    ws_sum.set_tab_color(Config.colors()['blue'])

    ws_sum.set_column("A:A", 7, format_align_center)
    ws_sum.set_column("B:B", 25, format_align_center)
    ws_sum.set_column("C:C", 24, format_align_center)
    ws_sum.set_column("D:D", 20, format_align_center)
    ws_sum.set_column("E:E", 7, format_align_center)

    # --------------------
    # VULN SUMMARY
    # --------------------
    ws_sum.merge_range("B2:D2", "VULNERABILITY SUMMARY", format_sheet_title_content)
    ws_sum.write("B3", "Threat", format_table_titles)
    ws_sum.write("C3", "Vulns number", format_table_titles)
    ws_sum.write("D3", "Affected hosts", format_table_titles)

    for i, level in enumerate(Config.levels().values(), 4):
        ws_sum.write("B{}".format(i), level.capitalize(), format_sheet_title_content)
        ws_sum.write("C{}".format(i), vuln_levels[level], format_align_border)
        ws_sum.write("D{}".format(i), vuln_host_by_level[level], format_align_border)

    ws_sum.write("B9", "Total", format_table_titles)
    ws_sum.write_formula("C9", "=SUM($C$4:$C$8)", format_table_titles)
    ws_sum.write_formula("D9", "=SUM($D$4:$D$8)", format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_summary = workbook.add_chart({'type': 'pie'})
    chart_vulns_summary.add_series({
        'name': 'vulnerability summary by affected hosts',
        'categories': '={}!B4:B8'.format(sheet_name),
        'values': '={}!D4:D8'.format(sheet_name),
        'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
        'points': [
            {'fill': {'color': Config.colors()['critical']}},
            {'fill': {'color': Config.colors()['high']}},
            {'fill': {'color': Config.colors()['medium']}},
            {'fill': {'color': Config.colors()['low']}},
            {'fill': {'color': Config.colors()['none']}},
        ],
    })
    chart_vulns_summary.set_title({'name': 'Vulnerability summary', 'overlay': False, 'name_font': {'name': 'Tahoma'}})
    chart_vulns_summary.set_size({'width': 500, 'height': 300})
    chart_vulns_summary.set_legend({'position': 'right', 'font': {'name': 'Tahoma'}})
    ws_sum.insert_chart("F2", chart_vulns_summary)

    # --------------------
    # VULN BY FAMILY
    # --------------------
    ws_sum.merge_range("B19:C19", "VULNERABILITIES BY FAMILY", format_sheet_title_content)
    ws_sum.write("B20", "family", format_table_titles)
    ws_sum.write("C20", "vulns number", format_table_titles)

    last = 21
    for i, (family, number) in enumerate(iter(vuln_by_family.items()), last):
        ws_sum.write("B{}".format(i), family, format_align_border)
        ws_sum.write("C{}".format(i), number, format_align_border)
        last = i

    ws_sum.write("B{}".format(str(last + 1)), "Total", format_table_titles)
    ws_sum.write_formula("C{}".format(str(last + 1)), "=SUM($C$21:$C${})".format(last), format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_by_family = workbook.add_chart({'type': 'pie'})
    chart_vulns_by_family.add_series({
        'name': 'vulnerability summary by family',
        'categories': '={}!B21:B{}'.format(sheet_name, last),
        'values': '={}!C21:C{}'.format(sheet_name, last),
        'data_labels': {'value': True, 'position': 'best_fit', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
    })
    chart_vulns_by_family.set_title({'name': 'Vulnerability by family', 'overlay': False,
                                     'name_font': {'name': 'Tahoma'}})
    chart_vulns_by_family.set_size({'width': 500, 'height': 500})
    chart_vulns_by_family.set_legend({'position': 'bottom', 'font': {'name': 'Tahoma'}})
    ws_sum.insert_chart("F19", chart_vulns_by_family)

    # ====================
    # TABLE OF CONTENTS
    # ====================
    sheet_name = "TOC"
    ws_toc = workbook.add_worksheet(sheet_name)
    ws_toc.set_tab_color(Config.colors()['blue'])

    ws_toc.set_column("A:A", 7)
    ws_toc.set_column("B:B", 5)
    ws_toc.set_column("C:C", 150)
    ws_toc.set_column("D:D", 15)
    ws_toc.set_column("E:E", 50)
    ws_toc.set_column("F:F", 7)

    ws_toc.merge_range("B2:E2", "TABLE OF CONTENTS", format_sheet_title_content)
    ws_toc.write("B3", "No.", format_table_titles)
    ws_toc.write("C3", "Vuln Title", format_table_titles)
    ws_toc.write("D3", "Level", format_table_titles)
    ws_toc.write("E3", "Hosts", format_table_titles)

    # ====================
    # VULN SHEETS
    # ====================
    for i, vuln in enumerate(vuln_info, 1):
        name = re.sub(r"[\[\]\\\'\"&@#():*?/]", "", vuln.name)
        if len(name) > 27:
            name = "{}..{}".format(name[0:15], name[-10:])
        name = "{:03X}_{}".format(i, name)
        ws_vuln = workbook.add_worksheet(name)
        ws_vuln.set_tab_color(Config.colors()[vuln.level.lower()])

        # --------------------
        # TABLE OF CONTENTS
        # --------------------
        ws_toc.write("B{}".format(i + 3), "{:03X}".format(i), format_table_cells)
        ws_toc.write_url("C{}".format(i + 3), "internal:'{}'!A1".format(name), format_table_cells, string=vuln.name)
        ws_toc.write("D{}".format(i + 3), "{:.1f} ({})".format(vuln.cvss, vuln.level.capitalize()),
                     format_toc[vuln.level])
        ws_toc.write("E{}".format(i + 3), "{}".format(', '.join([host.ip for host, _ in vuln.hosts])),
                     format_table_cells)
        ws_vuln.write_url("A1", "internal:'{}'!A{}".format(ws_toc.get_name(), i + 3), format_align_center,
                          string="<< TOC")
        ws_toc.set_row(i + 3, __row_height(name, 150), None)

        # --------------------
        # VULN INFO
        # --------------------
        ws_vuln.set_column("A:A", 7, format_align_center)
        ws_vuln.set_column("B:B", 20, format_align_center)
        ws_vuln.set_column("C:C", 20, format_align_center)
        ws_vuln.set_column("D:D", 50, format_align_center)
        ws_vuln.set_column("E:E", 15, format_align_center)
        ws_vuln.set_column("F:F", 15, format_align_center)
        ws_vuln.set_column("G:G", 20, format_align_center)
        ws_vuln.set_column("H:H", 7, format_align_center)
        content_width = 120

        ws_vuln.write('B2', "Title", format_table_titles)
        ws_vuln.merge_range("C2:G2", vuln.name, format_sheet_title_content)
        ws_vuln.set_row(1, __row_height(vuln.name, content_width), None)

        ws_vuln.write('B3', "Description", format_table_titles)
        ws_vuln.merge_range("C3:G3", vuln.description, format_table_cells)
        ws_vuln.set_row(2, __row_height(vuln.description, content_width), None)

        ws_vuln.write('B4', "Impact", format_table_titles)
        ws_vuln.merge_range("C4:G4", vuln.impact, format_table_cells)
        ws_vuln.set_row(3, __row_height(vuln.impact, content_width), None)

        ws_vuln.write('B5', "Recommendation", format_table_titles)
        ws_vuln.merge_range("C5:G5", vuln.solution, format_table_cells)
        ws_vuln.set_row(4, __row_height(vuln.solution, content_width), None)

        ws_vuln.write('B6', "Details", format_table_titles)
        ws_vuln.merge_range("C6:G6", vuln.insight, format_table_cells)
        ws_vuln.set_row(5, __row_height(vuln.insight, content_width), None)

        ws_vuln.write('B7', "CVEs", format_table_titles)
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"
        ws_vuln.merge_range("C7:G7", cves, format_table_cells)
        ws_vuln.set_row(6, __row_height(cves, content_width), None)

        ws_vuln.write('B8', "CVSS", format_table_titles)
        cvss = vuln.cvss if vuln.cvss != -1.0 else "No CVSS"
        ws_vuln.merge_range("C8:G8", "{:.1f}".format(cvss), format_table_cells)

        ws_vuln.write('B9', "Level", format_table_titles)
        ws_vuln.merge_range("C9:G9", vuln.level.capitalize(), format_table_cells)

        ws_vuln.write('B10', "Family", format_table_titles)
        ws_vuln.merge_range("C10:G10", vuln.family, format_table_cells)

        ws_vuln.write('C12', "IP", format_table_titles)
        ws_vuln.write('D12', "Host name", format_table_titles)
        ws_vuln.write('E12', "Port number", format_table_titles)
        ws_vuln.write('F12', "Port protocol", format_table_titles)

        # --------------------
        # AFFECTED HOSTS
        # --------------------
        for j, (host, port) in enumerate(vuln.hosts, 13):

            ws_vuln.write("C{}".format(j), host.ip)
            ws_vuln.write("D{}".format(j), host.host_name if host.host_name else "-")

            if port:
                ws_vuln.write("E{}".format(j), "" if port.number == 0 else port.number)
                ws_vuln.write("F{}".format(j), port.protocol)
            else:
                ws_vuln.write("E{}".format(j), "No port info")

    workbook.close()


def export_to_word(vuln_info, template, output_file='openvas_report.docx'):
    """
    Export vulnerabilities info in a Word file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param output_file: Filename of the Excel file
    :type output_file: str
    
    :param template: Path to Docx template
    :type template: str

    :raises: TypeError
    """

    import matplotlib.pyplot as plt
    import numpy as np
    import tempfile
    import os

    from docx import Document
    from docx.oxml.shared import qn, OxmlElement
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Cm

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        if not isinstance(template, str):
            raise TypeError("Expected str, got '{}' instead".format(type(template)))
    else:
        template = 'openvasreporting/src/openvas-template.docx'

    # TODO Move to function to de-duplicate this
    vuln_info.sort(key=lambda key: key.cvss, reverse=True)
    vuln_levels = Counter()
    vuln_host_by_level = Counter()
    vuln_by_family = Counter()

    for i, vuln in enumerate(vuln_info, 1):
        vuln_levels[vuln.level.lower()] += 1
        vuln_host_by_level[vuln.level.lower()] += len(vuln.hosts)
        vuln_by_family[vuln.family] += 1

    # ====================
    # DOCUMENT PROPERTIES
    # ====================
    document = Document(template)

    doc_prop = document.core_properties
    doc_prop.title = "OpenVAS Report"
    doc_prop.category = "Report"

    document.add_paragraph('OpenVAS Report', style='Title')

    # ====================
    # TABLE OF CONTENTS
    # ====================
    document.add_paragraph('Table of Contents', style='Heading 1')

    par = document.add_paragraph()
    run = par.add_run()
    fld_char = OxmlElement('w:fldChar')  # creates a new element
    fld_char.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instr_text.text = r'TOC \h \z \t "OV-H1toc;1;OV-H2toc;2;OV-H3toc;3;OV-Finding;3"'

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:t')
    fld_char3.text = "# Right-click to update field. #"
    fld_char2.append(fld_char3)

    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)

    document.add_page_break()

    # ====================
    # MANAGEMENT SUMMARY
    # ====================
    document.add_paragraph('Management Summary', style='OV-H1toc')
    document.add_paragraph('< TYPE YOUR MANAGEMENT SUMMARY HERE >')
    document.add_page_break()

    # ====================
    # TECHNICAL FINDINGS
    # ====================
    document.add_paragraph('Technical Findings', style='OV-H1toc')
    document.add_paragraph('The section below discusses the technical findings.')

    # --------------------
    # SUMMARY TABLE
    # --------------------
    document.add_paragraph('Summary', style='OV-H2toc')

    colors_sum = []
    labels_sum = []
    vuln_sum = []
    aff_sum = []

    table_summary = document.add_table(rows=1, cols=3)
    hdr_cells = table_summary.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Risk level').bold = True
    hdr_cells[1].paragraphs[0].add_run('Vulns number').bold = True
    hdr_cells[2].paragraphs[0].add_run('Affected hosts').bold = True

    # Provide data to table and charts
    for level in Config.levels().values():
        row_cells = table_summary.add_row().cells
        row_cells[0].text = level.capitalize()
        row_cells[1].text = str(vuln_levels[level])
        row_cells[2].text = str(vuln_host_by_level[level])
        colors_sum.append(Config.colors()[level])
        labels_sum.append(level)
        vuln_sum.append(vuln_levels[level])
        aff_sum.append(vuln_host_by_level[level])

    # --------------------
    # CHART
    # --------------------
    fd, path = tempfile.mkstemp(suffix='.png')

    par_chart = document.add_paragraph()
    run_chart = par_chart.add_run()

    plt.figure()

    pos = np.arange(len(labels_sum))
    width = 0.35

    bars_vuln = plt.bar(pos - width / 2, vuln_sum, width, align='center', label='Vulnerabilities',
                        color=colors_sum, edgecolor='black')
    bars_aff = plt.bar(pos + width / 2, aff_sum, width, align='center', label='Affected hosts',
                       color=colors_sum, edgecolor='black', hatch='//')
    plt.title('Vulnerability summary by risk level')
    plt.subplot().set_xticks(pos)
    plt.subplot().set_xticklabels(labels_sum)
    plt.gca().spines['left'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['bottom'].set_position('zero')
    plt.tick_params(top=False, bottom=True, left=False, right=False,
                    labelleft=False, labelbottom=True)
    plt.subplots_adjust(left=0.0, right=1.0)

    def __label_bars(barcontainer):
        for bar in barcontainer:
            height = bar.get_height()
            plt.gca().text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, str(int(height)),
                           ha='center', color='black', fontsize=9)

    __label_bars(bars_vuln)
    __label_bars(bars_aff)

    plt.legend()

    plt.savefig(path)

    # plt.show()  # DEBUG

    run_chart.add_picture(path, width=Cm(8.0))
    os.remove(path)

    plt.figure()

    values = list(vuln_by_family.values())
    pie, tx, autotexts = plt.pie(values, labels=vuln_by_family.keys(), autopct='')
    plt.title('Vulnerability by family')
    for i, txt in enumerate(autotexts):
        txt.set_text('{}'.format(values[i]))
    plt.axis('equal')

    plt.savefig(path, bbox_inches='tight')  # bbox_inches fixes labels being cut, however only on save not on show

    # plt.show()  # DEBUG

    run_chart.add_picture(path, width=Cm(8.0))
    os.remove(path)

    # ====================
    # VULN PAGES
    # ====================
    cur_level = ""

    for i, vuln in enumerate(vuln_info, 1):
        # --------------------
        # GENERAL
        # --------------------
        level = vuln.level.lower()

        if level != cur_level:
            document.add_paragraph(
                level.capitalize(), style='OV-H2toc').paragraph_format.page_break_before = True
            cur_level = level
        else:
            document.add_page_break()

        title = "[{}] {}".format(level.upper(), vuln.name)
        document.add_paragraph(title, style='OV-Finding')

        table_vuln = document.add_table(rows=7, cols=3)
        table_vuln.autofit = False

        # COLOR
        # --------------------
        col_cells = table_vuln.columns[0].cells
        col_cells[0].merge(col_cells[6])
        color_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), Config.colors()[vuln.level][1:]))
        col_cells[0]._tc.get_or_add_tcPr().append(color_fill)

        for col_cell in col_cells:
            col_cell.width = Cm(0.42)

        # TABLE HEADERS
        # --------------------
        hdr_cells = table_vuln.columns[1].cells
        hdr_cells[0].paragraphs[0].add_run('Description').bold = True
        hdr_cells[1].paragraphs[0].add_run('Impact').bold = True
        hdr_cells[2].paragraphs[0].add_run('Recommendation').bold = True
        hdr_cells[3].paragraphs[0].add_run('Details').bold = True
        hdr_cells[4].paragraphs[0].add_run('CVSS').bold = True
        hdr_cells[5].paragraphs[0].add_run('CVEs').bold = True
        hdr_cells[6].paragraphs[0].add_run('Family').bold = True

        for hdr_cell in hdr_cells:
            hdr_cell.width = Cm(3.58)

        # FIELDS
        # --------------------
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"

        cvss = str(vuln.cvss) if vuln.cvss != -1.0 else "No CVSS"

        txt_cells = table_vuln.columns[2].cells
        txt_cells[0].text = vuln.description
        txt_cells[1].text = vuln.impact
        txt_cells[2].text = vuln.solution
        txt_cells[3].text = vuln.insight
        txt_cells[4].text = cvss
        txt_cells[5].text = cves
        txt_cells[6].text = vuln.family

        for txt_cell in txt_cells:
            txt_cell.width = Cm(12.50)

        # VULN HOSTS
        # --------------------
        document.add_paragraph('Vulnerable hosts', style='Heading 4')

        table_hosts = document.add_table(cols=4, rows=(len(vuln.hosts) + 1))
        hdr_cells = table_hosts.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('IP').bold = True
        hdr_cells[1].paragraphs[0].add_run('Host name').bold = True
        hdr_cells[2].paragraphs[0].add_run('Port number').bold = True
        hdr_cells[3].paragraphs[0].add_run('Port protocol').bold = True

        for j, (host, port) in enumerate(vuln.hosts, 1):

            cells = table_hosts.rows[j].cells
            cells[0].text = host.ip
            cells[1].text = host.host_name if host.host_name else "-"
            if port and port is not None:
                cells[2].text = "-" if port.number == 0 else str(port.number)
                cells[3].text = port.protocol
            else:
                cells[2].text = "No port info"

    document.save(output_file)


def export_to_csv(vuln_info, template=None, output_file='openvas_report.csv'):
    """
    Export vulnerabilities info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in CSV-output.")

    # TODO Move to function to de-duplicate this
    vuln_info.sort(key=lambda key: key.cvss, reverse=True)

    with open(output_file, 'w') as csvfile:
        fieldnames = ['hostname', 'ip', 'port', 'protocol',
                      'vulnerability', 'cvss', 'threat', 'family',
                      'description', 'detection', 'insight', 'impact', 'affected', 'solution', 'solution_type',
                      'vuln_id', 'cve', 'references']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for vuln in vuln_info:
            for (host, port) in vuln.hosts:
                rowdata = {
                    'hostname': host.host_name,
                    'ip': host.ip,
                    'port': port.number,
                    'protocol': port.protocol,
                    'vulnerability': vuln.name,
                    'cvss': vuln.cvss,
                    'threat': vuln.level,
                    'family': vuln.family,
                    'description': vuln.description,
                    'detection': vuln.detect,
                    'insight': vuln.insight,
                    'impact': vuln.impact,
                    'affected': vuln.affected,
                    'solution': vuln.solution,
                    'solution_type': vuln.solution_type,
                    'vuln_id': vuln.vuln_id,
                    'cve': ' - '.join(vuln.cves),
                    'references': ' - '.join(vuln.references)
                }
                writer.writerow(rowdata)
