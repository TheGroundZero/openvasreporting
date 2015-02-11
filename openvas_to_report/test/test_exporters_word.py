# -*- coding: utf-8 -*-
#
#
# Project name: OpenVAS2Report: A set of tools to manager OpenVAS XML report files.
# Project URL: https://github.com/cr0hn/openvas_to_report
#
# Copyright (c) 2015, cr0hn<-AT->cr0hn.com
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without modification, are permitted provided that the
# following conditions are met:
#
# 1. Redistributions of source code must retain the above copyright notice, this list of conditions and the
# following disclaimer.
#
# 2. Redistributions in binary form must reproduce the above copyright notice, this list of conditions and the
# following disclaimer in the documentation and/or other materials provided with the distribution.
#
# 3. Neither the name of the copyright holder nor the names of its contributors may be used to endorse or promote
# products derived from this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND ANY EXPRESS OR IMPLIED WARRANTIES,
# INCLUDING, BUT NOT LIMITED TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL,
# SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY,
# WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#

from docx.table import Table, _Cell

import pytest
import os

from docx import Document
from docx.text import Paragraph

from ..libs.exporters.word import export_to_word, create_tables, search_tables_in_document, \
    search_in_table, paragraph_search
from ..libs.data.parsed_data import Vulnerability, Host, Port

W_PATH = os.path.abspath(os.path.join(os.getcwd(), "test_word.docx"))
W_TEMPLATE = os.path.abspath(os.path.join(os.getcwd(), "template.docx"))


# ----------------------------------------------------------------------
@pytest.fixture
def generate_vuln():
    v = Vulnerability("111", "vuln name", "Low")
    v.add_host(Host("ip"), Port(80))
    return [v]


# --------------------------------------------------------------------------
# export_to_word test
# --------------------------------------------------------------------------
class TestExportToWord:

    # ----------------------------------------------------------------------
    def test_types(self, generate_vuln):
        pytest.raises(TypeError, export_to_word, None, "")
        pytest.raises(TypeError, export_to_word, 0, "")
        pytest.raises(TypeError, export_to_word, [""], None)
        pytest.raises(TypeError, export_to_word, generate_vuln, None)
        pytest.raises(ValueError, export_to_word, generate_vuln, "")
        pytest.raises(TypeError, export_to_word, generate_vuln, W_PATH, 0)
        pytest.raises(ValueError, export_to_word, generate_vuln, W_PATH, "")

    # ----------------------------------------------------------------------
    def test_if_file_is_created(self, generate_vuln):

        # Info for the word
        v = generate_vuln

        # Create the file
        export_to_word(v, W_PATH)

        # Checks if file exits
        assert os.path.exists(W_PATH)

        # If exits, delete tmp file
        os.remove(W_PATH)


# --------------------------------------------------------------------------
# export_to_word test
# --------------------------------------------------------------------------
class TestCreateTable:

    # ----------------------------------------------------------------------
    def test_types(self, generate_vuln):
        doc = Document()

        pytest.raises(TypeError, create_tables, None, doc)
        pytest.raises(TypeError, create_tables,  generate_vuln, None)

    # ----------------------------------------------------------------------
    def test_basic_pass_paragraph(self, generate_vuln):

        doc = Document()
        paragraph = doc.add_paragraph("my_paragraph")

        create_tables(generate_vuln, paragraph)

    # ----------------------------------------------------------------------
    def test_modify_paragraph(self, generate_vuln):
        pass
        #
        # doc = Document()
        # paragraph = doc.add_paragraph("my_paragraph")
        #
        # paragraph = create_tables(generate_vuln, paragraph)
        #
        # # Check results
        # assert doc.paragraphs[0].text == paragraph.text


# --------------------------------------------------------------------------
# search_tables_in_document test
# --------------------------------------------------------------------------
class TestSearchTablesInDocument:

    # ----------------------------------------------------------------------
    def test_types(self, generate_vuln):
        doc = Document()

        pytest.raises(TypeError, search_tables_in_document, None, doc)
        pytest.raises(TypeError, search_tables_in_document,  "", None)

    # ----------------------------------------------------------------------
    def test_get_table_by_key(self):

        doc = Document(W_TEMPLATE)

        t = search_tables_in_document("####VULN_ID####", doc)

        assert isinstance(t, Table)

    # ----------------------------------------------------------------------
    def test_get_table_by_key_invalid_regex(self):

        doc = Document(W_TEMPLATE)

        pytest.raises(ValueError, search_tables_in_document, "????????", doc)

    # ----------------------------------------------------------------------
    def test_get_table_by_key_not_found(self):

        doc = Document(W_TEMPLATE)

        t = search_tables_in_document("===================", doc)

        assert t is None


# --------------------------------------------------------------------------
# search_in_table test
# --------------------------------------------------------------------------
class TestSearchInTable:

    # ----------------------------------------------------------------------
    def test_types(self, generate_vuln):
        doc = Document()

        pytest.raises(TypeError, search_in_table, None, doc)
        pytest.raises(TypeError, search_in_table,  "", None)

    # ----------------------------------------------------------------------
    def test_get_value_in_table(self):

        doc = Document(W_TEMPLATE)

        # Get table
        t = search_tables_in_document("####VULN_ID####", doc)

        # Search in talbe
        text_in_table = search_in_table("Vulnerabilidad", t)

        assert isinstance(text_in_table, _Cell)

        # Try to get texts
        texts = "".join(x.text for x in text_in_table.paragraphs)
        assert "Vulnerabilidad" in texts

    # ----------------------------------------------------------------------
    def test_get_table_by_key_invalid_regex(self):

        doc = Document(W_TEMPLATE)
        t = search_tables_in_document("####VULN_ID####", doc)

        pytest.raises(ValueError, search_in_table, "????????", t)

    # ----------------------------------------------------------------------
    def test_get_table_by_key_not_found(self):

        doc = Document(W_TEMPLATE)
        t = search_tables_in_document("####VULN_ID####", doc)

        t2 = search_in_table("===================", t)

        assert t2 is None


# --------------------------------------------------------------------------
# paragraph_search test
# --------------------------------------------------------------------------
class TestParagraphSearch:

    # ----------------------------------------------------------------------
    def test_types(self,):
        doc = Document()

        pytest.raises(TypeError, paragraph_search, None, doc)
        pytest.raises(TypeError, paragraph_search,  "", None)

    # ----------------------------------------------------------------------
    def test_text(self):

        doc = Document(W_TEMPLATE)

        t = paragraph_search("#####TABLE####", doc)

        assert isinstance(t, Paragraph)

    # ----------------------------------------------------------------------
    def test_get_table_by_key_invalid_regex(self):

        doc = Document(W_TEMPLATE)

        pytest.raises(ValueError, paragraph_search, "????????", doc)

    # ----------------------------------------------------------------------
    def test_get_table_by_key_not_found(self):

        doc = Document(W_TEMPLATE)

        t = paragraph_search("===================", doc)

        assert t is None
