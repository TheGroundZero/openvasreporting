# -*- coding: utf-8 -*-
#
#
# Project name: OpenVAS2Report: A set of tools to manager OpenVAS XML report files.
# Project URL: https://github.com/cr0hn/openvas_to_report
#
# Copyright (c) 2015, cr0hn<-AT->cr0hn.com
# All rights reserved.
#
# Redistribution and use in source and binary forms, with or without modification, are permitted provided that the
# following conditions are met:
#
# 1. Redistributions of source code must retain the above copyright notice, this list of conditions and the
# following disclaimer.
#
# 2. Redistributions in binary form must reproduce the above copyright notice, this list of conditions and the
# following disclaimer in the documentation and/or other materials provided with the distribution.
#
# 3. Neither the name of the copyright holder nor the names of its contributors may be used to endorse or promote
# products derived from this software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS" AND ANY EXPRESS OR IMPLIED WARRANTIES,
# INCLUDING, BUT NOT LIMITED TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL,
# SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY,
# WHETHER IN CONTRACT, STRICT LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#


__all__ = ["export_to_word"]

import re
import os

from docx import Document
from docx.text import Paragraph, Run
from docx.table import Table

from ..data.parsed_data import Vulnerability

#
# Search and replace: https://github.com/sk1tt1sh/python-docx/blob/develop/docx/api.py
#


# ----------------------------------------------------------------------
def export_to_word(vuln_info, output_file_name, template=None):
    """
    Export Vulnerability info to word file.
    
    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)
    
    :param output_file_name: filename to save word file,
    :type output_file_name: str

    :param template: path to template file.
    :type template: str

    :raises: TypeError
    """
    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '%s' instead" % type(vuln_info))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '%s' instead" % type(x))
    if not isinstance(output_file_name, str):
        raise TypeError("Expected basestring, got '%s' instead" % type(output_file_name))
    else:
        if not output_file_name:
            raise ValueError("output_file_name must has a valid value.")
    if template is not None:
        if not isinstance(template, str):
            raise TypeError("Expected basestring, got '%s' instead" % type(template))
        if not template or not os.path.exists(template):
            raise ValueError("Template must has a valid value.")

    if template:
        doc_template = Document(template)

        # Get mark where insert tables
        p = paragraph_search("#####TABLE####", doc_template)
        if p is None:
            raise IOError("Incompatible document format. Can't find '#####TABLE####' in document.")
    else:
        p = Paragraph("")

    # Add table information
    create_tables(vuln_info, p)

    # Save doc
    doc = Document()
    # doc.add_paragraph(p)
    doc.save(output_file_name)


# ----------------------------------------------------------------------
def create_tables(vulns, paragraph):
    """
    Creates tables with information vulns into paragraph instance.
    
    :param vulns: list of Vulnerabilities instances 
    :type vulns: list(Vulnerability)
    
    :param paragraph: A Paragraph instance.
    :type paragraph: Paragraph

    :raises: TypeError
    """
    if not isinstance(vulns, list):
        raise TypeError("Expected list, got '%s' instead" % type(vulns))
    if not isinstance(paragraph, Paragraph):
        raise TypeError("Expected Paragraph, got '%s' instead" % type(paragraph))

    d = Document()

    return d.add_paragraph("asdfasdfasfd")


# ----------------------------------------------------------------------
def paragraph_search(text_value, doc):
    """
    Search in paragraphs for the text.

    :param text_value: Text to looking for in document.
    :type text_value: str

    :param doc: Document instance
    :type doc: Document

    :return: paragraph found. None otherwise.
    :rtype: Paragraph|None

    :raises: ValueError, TypeError
    """
    if not isinstance(text_value, str):
        raise TypeError("Expected basestring, got '%s' instead" % type(text_value))
    if not isinstance(doc, Document):
        raise TypeError("Expected Document, got '%s' instead" % type(doc))

    try:
        para_regex = re.compile(text_value)
    except re.error:
        raise ValueError("Invalid regex expression.")

    for paragraph in doc.paragraphs:
        if paragraph.text:
            if para_regex.search(paragraph.text):
                return paragraph
    return None


# ----------------------------------------------------------------------
def search_tables_in_document(text_value, doc):
    """
    Search for a table into document and return the cell where found it.

    :param text_value: Value to search, as regex.
    :type text_value: str

    :param doc: Document instance
    :type doc: Document

    :return: A Table instances or None, if not matches found.
    :rtype: Table

    :raises: ValueError, TypeError
    """
    if not isinstance(text_value, str):
        raise TypeError("Expected basestring, got '%s' instead" % type(text_value))
    if not isinstance(doc, Document):
        raise TypeError("Expected Document, got '%s' instead" % type(doc))

    try:
        tbl_regex = re.compile(text_value)
    except re.error:
        raise ValueError("Invalid regex expression.")

    if tbl_regex:
        for table in doc.tables:
            for r in table.rows:
                for cell in r.cells:
                    for paragraph in cell.paragraphs:
                        print(paragraph.text)
                        if paragraph.text:
                            if tbl_regex.search(paragraph.text):
                                return table
    return None


# ----------------------------------------------------------------------
def search_in_table(text_value, table):
    """
    Search for a text into a table and return the cell where found it.

    :param text_value: Value to search as regex.
    :type text_value: str

    :param table: Table instance
    :type table: Table

    :return: A _Cell instances or None, if not matches found.
    :rtype: _Cell

    :raises: TypeError, ValueError
    """
    if not isinstance(text_value, str):
        raise TypeError("Expected basestring, got '%s' instead" % type(text_value))
    if not isinstance(table, Table):
        raise TypeError("Expected Document, got '%s' instead" % type(table))

    try:
        tbl_regex = re.compile(text_value)
    except re.error:
        raise ValueError("Invalid regex expression.")

    for r in table.rows:
        for cell in r.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text:
                    if tbl_regex.search(paragraph.text):
                        return cell
    return None